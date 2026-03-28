"""
Generate SDE_Spark Implementation Presentation (.docx)
A detailed document covering the full implementation, architecture,
testing, and experimental evaluation of the SDE_Spark project.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Style Setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Dark blue

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(3)
code_style.paragraph_format.space_after = Pt(3)
code_style.paragraph_format.left_indent = Cm(1)

def add_code(text, doc=doc):
    """Add a code block."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line, style='CodeBlock')

def add_table(headers, rows, doc=doc):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table


# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SDE_Spark')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Synopsis Data Engine\nApache Spark Implementation')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Detailed Implementation Report')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    'Based on the SDE Architecture by Kontaxakis (2020)\n'
    'Re-implemented on Apache Spark 3.5.3 Structured Streaming\n\n'
    'March 2026'
)
run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Technology Stack & Project Structure',
    '3. Message Model & Kafka Contract',
    '4. Pipeline Architecture Overview',
    '5. Layer 1 — Ingestion & Parsing',
    '6. Layer 2 — Routing & Registration',
    '7. Layer 3 — Synopsis Processing (Core)',
    '8. Layer 4 — Path Splitting',
    '9. Layer 5 — Aggregation (Purple Path)',
    '10. Layer 6 — Output',
    '11. Synopsis Algorithms',
    '12. Configuration & Deployment',
    '13. State Management & Serialization',
    '14. Monitoring & Metrics',
    '15. End-to-End Validated Flows',
    '16. Unit Testing',
    '17. Experimental Evaluation',
    '18. Comparison with SDE_Flink',
    '19. Conclusions & Future Work',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'The Synopsis Data Engine (SDE) is a distributed system for maintaining and querying '
    'probabilistic data structures (synopses) over continuous data streams. The original '
    'SDE was implemented on Apache Flink 1.9.3 as part of the INFORE EU research project '
    '(Kontaxakis, 2020). This document presents SDE_Spark, a complete re-implementation '
    'of the SDE on Apache Spark 3.5.3 Structured Streaming.'
)

doc.add_paragraph(
    'The re-implementation preserves the original 6-layer pipeline architecture and the '
    'Kafka message contract, ensuring full compatibility with existing upstream and downstream '
    'systems. At the same time, it leverages Spark-specific features such as flatMapGroupsWithState '
    'for stateful processing, Kryo serialization, and Spark\'s built-in checkpointing mechanism.'
)

doc.add_heading('1.1 Objectives', level=2)
objectives = [
    'Preserve the SDE architecture and Kafka API contract from the Flink implementation',
    'Implement all 4 synopsis algorithms: CountMin Sketch, Bloom Filter, AMS Sketch, HyperLogLog',
    'Support both GREEN path (noOfP=1, single-partition) and PURPLE path (noOfP>1, parallel processing with aggregation)',
    'Provide production-ready configuration, monitoring, and deployment capabilities',
    'Validate correctness through unit tests and end-to-end integration tests',
    'Evaluate performance through experiments comparable to Kontaxakis (2020) Chapter 5',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph(
    'This document covers the complete implementation of SDE_Spark: the pipeline architecture, '
    'each processing layer in detail, synopsis algorithms, configuration system, state management, '
    'testing strategy, and preliminary experimental evaluation. It serves as both a technical '
    'reference and a thesis chapter documenting the implementation work.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════════

doc.add_heading('2. Technology Stack & Project Structure', level=1)

doc.add_heading('2.1 Technology Stack', level=2)

add_table(
    ['Component', 'Technology', 'Version'],
    [
        ['Stream Processing', 'Apache Spark Structured Streaming', '3.5.3'],
        ['Message Broker', 'Apache Kafka', '7.5.3'],
        ['Language', 'Java', '17'],
        ['Build System', 'Apache Maven', '3.x'],
        ['Serialization', 'Kryo (via Spark)', '4.x'],
        ['JSON Processing', 'Jackson Databind', '2.15.3'],
        ['Probabilistic DS', 'stream-lib (Clearspring)', '2.9.5'],
        ['Streaming Algorithms', 'streaminer', '1.1.1'],
        ['Logging', 'SLF4J + Logback', '2.0.9 / 1.4.14'],
        ['Testing', 'JUnit 5 + AssertJ', '5.10.1 / 3.24.2'],
        ['Packaging', 'Maven Shade Plugin (fat JAR)', '3.5.1'],
    ]
)

doc.add_heading('2.2 Project Structure', level=2)

doc.add_paragraph(
    'SDE_Spark is organized as a standard Maven project with the following package structure:'
)

add_code("""SDE_Spark/
├── pom.xml
├── docker-compose.yml                    # Kafka + Zookeeper
├── src/main/java/infore/sde/spark/
│   ├── SDESparkApp.java                  # Main entry point
│   ├── config/
│   │   └── SDEConfig.java                # Centralized configuration
│   ├── ingestion/
│   │   └── KafkaIngestionLayer.java      # Layer 1
│   ├── routing/
│   │   ├── DataRouter.java               # Layer 2 (stateful)
│   │   ├── RequestRouter.java            # Request fan-out (standalone)
│   │   └── RoutingState.java             # Routing state model
│   ├── processing/
│   │   ├── InputEvent.java               # Union discriminator
│   │   ├── SynopsisProcessor.java        # Layer 3 (stateful)
│   │   └── SynopsisProcessorState.java   # Synopsis state model
│   ├── aggregation/
│   │   ├── PathSplitter.java             # Layer 4
│   │   ├── ReduceAggregator.java         # Layer 5 (stateful)
│   │   └── AggregationState.java         # Aggregation state model
│   ├── output/
│   │   └── KafkaOutputLayer.java         # Layer 6
│   ├── messages/
│   │   ├── Datapoint.java                # Data message POJO
│   │   ├── Request.java                  # Request message POJO
│   │   └── Estimation.java               # Output message POJO
│   ├── synopses/
│   │   ├── Synopsis.java                 # Abstract base class
│   │   ├── SynopsisFactory.java          # Factory pattern
│   │   ├── CountMin.java                 # CountMin Sketch
│   │   ├── Bloomfilter.java              # Bloom Filter
│   │   ├── AMSsynopsis.java              # AMS Sketch
│   │   ├── HyperLogLogSynopsis.java      # HyperLogLog
│   │   └── sketches/CM.java              # CountMin core impl
│   ├── reduceFunctions/
│   │   ├── ReduceFunction.java           # Abstract reducer
│   │   ├── SimpleSumFunction.java        # Sum aggregation
│   │   └── SimpleORFunction.java         # OR aggregation
│   ├── metrics/
│   │   ├── PipelineMetrics.java          # Spark accumulators
│   │   └── ThroughputListener.java       # Per-batch CSV metrics
│   └── integration/
│       ├── LoadTestProducer.java          # Experiment load generator
│       ├── EndToEndTestApp.java           # E2E test pipeline
│       ├── EndToEndTestProducer.java      # E2E test data producer
│       └── ...                            # Layer-specific test apps
├── src/test/java/                         # 12 unit test classes
├── src/main/resources/
│   ├── logback.xml
│   ├── log4j2.properties
│   └── sde-spark.properties.example
├── results/                               # Experiment outputs
└── docs/                                  # Documentation""")

doc.add_paragraph(
    'The project comprises 38 main Java source files and 12 test classes, '
    'organized into 12 packages following the layered architecture.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 3. MESSAGE MODEL
# ════════════════════════════════════════════════════════════════════

doc.add_heading('3. Message Model & Kafka Contract', level=1)

doc.add_paragraph(
    'SDE_Spark preserves the exact Kafka message contract from the Flink implementation. '
    'This is critical for backward compatibility: any existing producer or consumer that '
    'works with SDE_Flink will work unchanged with SDE_Spark. Three message types flow '
    'through the system, each mapped to a dedicated Kafka topic.'
)

doc.add_heading('3.1 Datapoint (data_topic)', level=2)
doc.add_paragraph(
    'Represents a single observation from a data stream. The dataSetKey identifies the '
    'logical dataset (e.g., "Forex"), streamID identifies the specific stream within '
    'that dataset (e.g., "EURUSD"), and values contains the actual data as a JSON object.'
)
add_code("""{
  "dataSetkey": "Forex",
  "streamID":   "EURUSD",
  "values":     {"StockID": "AAPL", "price": "185.50"}
}""")

doc.add_paragraph(
    'Implementation: Datapoint.java uses Jackson @JsonProperty annotations for '
    'deserialization. The values field is stored as a Jackson JsonNode to support '
    'arbitrary nested structures without schema coupling.'
)

doc.add_heading('3.2 Request (request_topic)', level=2)
doc.add_paragraph(
    'Commands sent by clients to manage synopsis lifecycle. The requestID encodes '
    'the operation type (last digit): 1=ADD, 2=DELETE, 3=ESTIMATE. '
    'The noOfP field controls parallelism: 1 for GREEN path, >1 for PURPLE path.'
)
add_code("""{
  "dataSetkey":  "Forex",
  "requestID":   1,
  "synopsisID":  1,
  "uid":         10,
  "streamID":    "ALL",
  "param":       ["StockID", "price", "Queryable", "0.01", "0.99", "42"],
  "noOfP":       1
}""")

add_table(
    ['Field', 'Type', 'Description'],
    [
        ['dataSetkey', 'String', 'Logical dataset key for routing'],
        ['requestID', 'int', 'Operation: %10 -> 1=ADD, 2=DELETE, 3=ESTIMATE'],
        ['synopsisID', 'int', 'Algorithm: 1=CountMin, 2=Bloom, 3=AMS, 4=HLL'],
        ['uid', 'int', 'Unique synopsis instance identifier'],
        ['streamID', 'String', 'Target stream (or "ALL")'],
        ['param', 'String[]', 'Algorithm-specific parameters'],
        ['noOfP', 'int', 'Parallelism level (1=GREEN, >1=PURPLE)'],
    ]
)

doc.add_heading('3.3 Estimation (estimation_topic)', level=2)
doc.add_paragraph(
    'The output message containing query results. Produced by the pipeline when '
    'an ESTIMATE request is processed.'
)
add_code("""{
  "key":            "Forex",
  "estimationkey":  "10",
  "uid":            10,
  "requestID":      3,
  "synopsisID":     1,
  "estimation":     "549.0",
  "param":          ["AAPL"],
  "noOfP":          1
}""")

doc.add_heading('3.4 InputEvent — The Union Discriminator', level=2)
doc.add_paragraph(
    'A key design challenge in Spark Structured Streaming is that flatMapGroupsWithState '
    'operates on a single typed Dataset. Since the SDE must process both data and requests '
    'in the same stateful operator, both streams are wrapped in an InputEvent discriminator:'
)
add_code("""public class InputEvent implements Serializable {
    public enum Type { DATA, REQUEST }

    private final Type type;
    private final Datapoint datapoint;
    private final Request request;
    private final String dataSetKey;

    public static InputEvent data(Datapoint dp) { ... }
    public static InputEvent request(Request rq) { ... }
}""")

doc.add_paragraph(
    'This pattern replaces Flink\'s ConnectedStreams, which natively supports '
    'two-input operators via CoProcessFunction.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 4. PIPELINE ARCHITECTURE OVERVIEW
# ════════════════════════════════════════════════════════════════════

doc.add_heading('4. Pipeline Architecture Overview', level=1)

doc.add_paragraph(
    'SDE_Spark implements the same 6-layer pipeline architecture as the original '
    'SDE_Flink. The pipeline is constructed as a single Spark Structured Streaming '
    'query that chains together stateful and stateless transformations.'
)

add_table(
    ['Layer', 'Component', 'Spark API', 'Flink Equivalent'],
    [
        ['1. Ingestion', 'KafkaIngestionLayer', 'readStream().format("kafka")', 'FlinkKafkaConsumer'],
        ['2. Routing', 'DataRouter', 'flatMapGroupsWithState', 'CoProcessFunction'],
        ['3. Processing', 'SynopsisProcessor', 'flatMapGroupsWithState', 'CoProcessFunction'],
        ['4. Splitting', 'PathSplitter', 'filter()', 'SplitStream / OutputSelector'],
        ['5. Aggregation', 'ReduceAggregator', 'flatMapGroupsWithState', 'ReduceFlatMap'],
        ['6. Output', 'KafkaOutputLayer', 'writeStream().format("kafka")', 'FlinkKafkaProducer'],
    ]
)

doc.add_heading('4.1 Data Flow', level=2)
doc.add_paragraph(
    'The pipeline processes two input streams (data and requests) and produces '
    'one output stream (estimations). The flow follows two paths based on the '
    'parallelism level (noOfP):'
)

doc.add_paragraph('GREEN Path (noOfP = 1):', style='List Bullet')
doc.add_paragraph(
    'Kafka -> Ingestion -> Routing (pass-through) -> Synopsis Processing -> '
    'Path Split -> Direct Output'
)
doc.add_paragraph('PURPLE Path (noOfP > 1):', style='List Bullet')
doc.add_paragraph(
    'Kafka -> Ingestion -> Routing (fan-out to N partitions) -> Synopsis Processing '
    '(N parallel instances) -> Path Split -> Aggregation (merge N partial results) -> Output'
)

doc.add_heading('4.2 Main Entry Point — SDESparkApp.java', level=2)
doc.add_paragraph(
    'The main class wires all 6 layers together in a single streaming query. '
    'It creates a SparkSession with Kryo serialization, registers the ThroughputListener '
    'for metrics, and constructs the pipeline DAG:'
)

add_code("""// Layer 1: Ingestion
Dataset<Datapoint> dataStream = ingestion.readDataStream();
Dataset<Request> requestStream = ingestion.readRequestStream();

// Layer 2: Union + Routing
Dataset<InputEvent> combined = taggedData.union(taggedRequests);
Dataset<InputEvent> routed = combined.groupByKey(InputEvent::getDataSetKey)
    .flatMapGroupsWithState(new DataRouter(config), ...);

// Layer 3: Synopsis Processing
Dataset<Estimation> estimations = routed.groupByKey(InputEvent::getDataSetKey)
    .flatMapGroupsWithState(new SynopsisProcessor(config, metrics), ...);

// Layer 4: Path Splitting
PathSplitter splitter = new PathSplitter(estimations);

// Layer 5: Aggregation (Purple path only)
Dataset<Estimation> aggregated = splitter.getPurplePath()
    .groupByKey(Estimation::getUid)
    .flatMapGroupsWithState(new ReduceAggregator(config), ...);

// Merge GREEN + PURPLE and output
Dataset<Estimation> all = splitter.getGreenPath().union(aggregated);
StreamingQuery query = outputLayer.write(all, "sde-output");""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 5. LAYER 1 — INGESTION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('5. Layer 1 — Ingestion & Parsing', level=1)

doc.add_paragraph(
    'KafkaIngestionLayer is responsible for consuming raw JSON messages from Kafka '
    'and deserializing them into typed Java objects (Datapoint and Request). It replaces '
    'the FlinkKafkaConsumer and inline MapFunction from the original Flink Run.java.'
)

doc.add_heading('5.1 Implementation', level=2)
doc.add_paragraph(
    'The layer creates two independent Kafka consumers, one for the data topic and '
    'one for the request topic. Each consumer is configured with its own consumer group ID '
    '(suffixed with "-data" and "-request") to allow independent offset tracking.'
)

add_code("""public Dataset<Datapoint> readDataStream() {
    DataStreamReader reader = spark.readStream()
        .format("kafka")
        .option("kafka.bootstrap.servers", config.getKafkaBrokers())
        .option("subscribe", config.getDataTopic())
        .option("startingOffsets", config.getStartingOffsets())
        .option("kafka.group.id", config.getKafkaGroupId() + "-data")
        .option("failOnDataLoss", String.valueOf(config.isFailOnDataLoss()));
    // Apply security options (SASL, SSL) if configured
    config.getKafkaSecurityOptions().forEach(reader::option);

    Dataset<Row> raw = reader.load()
        .selectExpr("CAST(value AS STRING) as json");

    return raw.map(row -> {
        ObjectMapper mapper = new ObjectMapper();
        return mapper.readValue(row.getString(0), Datapoint.class);
    }, Encoders.kryo(Datapoint.class))
    .filter(dp -> dp != null && dp.getDataSetKey() != null);
}""")

doc.add_heading('5.2 Key Design Decisions', level=2)
decisions = [
    ('Kryo encoding', 'All POJOs use Encoders.kryo() instead of Spark\'s built-in bean encoder. '
     'This avoids schema inference issues with Jackson\'s JsonNode and provides better '
     'performance for complex nested objects.'),
    ('Null filtering', 'Malformed messages that fail JSON parsing return null and are filtered '
     'out, preventing poison-pill messages from crashing the pipeline.'),
    ('Separate consumer groups', 'Data and request consumers use distinct group IDs so they '
     'track offsets independently. This prevents a slow request consumer from blocking data processing.'),
    ('Configurable offsets', 'The startingOffsets parameter (default: "earliest") can be set to '
     '"latest" for production deployments where historical replay is not needed.'),
]
for title, desc in decisions:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 6. LAYER 2 — ROUTING
# ════════════════════════════════════════════════════════════════════

doc.add_heading('6. Layer 2 — Routing & Registration', level=1)

doc.add_paragraph(
    'The routing layer is the most architecturally significant change from the Flink '
    'implementation. In SDE_Flink, routing was split across two separate operators: '
    'RqRouterFlatMap (request fan-out) and dataRouterCoFlatMap (data routing). '
    'In SDE_Spark, these are unified into a single stateful operator — DataRouter — '
    'because request registration and data routing must share the same partition state.'
)

doc.add_heading('6.1 The Key Insight', level=2)
doc.add_paragraph(
    'In Flink, ConnectedStreams allows two inputs to share state within a single operator. '
    'Spark has no equivalent. Instead, both streams are unioned into a single Dataset<InputEvent> '
    'and processed together in flatMapGroupsWithState. The DataRouter maintains a RoutingState '
    'that tracks which parallelism levels are registered for each dataSetKey.'
)

doc.add_heading('6.2 DataRouter Implementation', level=2)
doc.add_paragraph('The DataRouter handles three categories of events:')

doc.add_paragraph('ADD/DELETE Requests (operation 1, 2):', style='List Bullet')
doc.add_paragraph(
    'When a request with noOfP > 1 arrives, the DataRouter registers the parallelism '
    'level in its state and fans out the request to N keyed partition keys '
    '(e.g., "Forex_2_KEYED_0", "Forex_2_KEYED_1"). For noOfP = 1, requests pass through unchanged.'
)

doc.add_paragraph('ESTIMATE Requests (operation 3):', style='List Bullet')
doc.add_paragraph(
    'Similarly fanned out to all N partition keys for PURPLE path synopses, '
    'or passed through for GREEN path.'
)

doc.add_paragraph('Data Events:', style='List Bullet')
doc.add_paragraph(
    'For each registered parallelism level, data is hash-partitioned by streamID '
    'to the appropriate keyed partition: slot = abs(streamID.hashCode()) % parallelism. '
    'Data is also forwarded with the original key for noOfP=1 synopses.'
)

doc.add_heading('6.3 Event Ordering', level=2)
doc.add_paragraph(
    'Within a Spark micro-batch, event ordering is not guaranteed. The DataRouter sorts '
    'events into three groups before processing: ADD/DELETE requests first, then data events, '
    'then other requests (ESTIMATE). This ensures that synopsis registrations are applied '
    'before data routing occurs within the same batch.'
)

add_code("""// Sort events: ADD/DELETE requests first, then data, then ESTIMATE
List<InputEvent> addDeleteRequests = new ArrayList<>();
List<InputEvent> dataEvents = new ArrayList<>();
List<InputEvent> otherRequests = new ArrayList<>();

for (InputEvent event : addDeleteRequests) handleRequest(...);
for (InputEvent event : dataEvents) handleData(...);
for (InputEvent event : otherRequests) handleRequest(...);""")

doc.add_heading('6.4 RoutingState', level=2)
doc.add_paragraph('The routing state tracks three pieces of information:')
add_table(
    ['Field', 'Type', 'Purpose'],
    [
        ['keyedParallelism', 'Map<Integer, Integer>', 'parallelism level -> count of synopses using it'],
        ['keysPerStream', 'Map<String, List<String>>', 'streamID -> routing keys (for future optimization)'],
        ['registrations', 'Map<Integer, Registration>', 'uid -> original registration (for DELETE cleanup)'],
    ]
)

doc.add_heading('6.5 State Timeout', level=2)
doc.add_paragraph(
    'The DataRouter uses ProcessingTimeTimeout to automatically evict stale routing state '
    'for keys that have no active synopses. The timeout duration is configurable via '
    '--routing-state-ttl (default: 1 day). When the state times out, all routing registrations '
    'for that key are removed.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 7. LAYER 3 — SYNOPSIS PROCESSING
# ════════════════════════════════════════════════════════════════════

doc.add_heading('7. Layer 3 — Synopsis Processing (Core)', level=1)

doc.add_paragraph(
    'The SynopsisProcessor is the heart of the SDE pipeline. It manages the lifecycle '
    'of synopsis data structures: creating them on ADD requests, feeding them data, '
    'querying them on ESTIMATE requests, and removing them on DELETE. It is implemented '
    'as a flatMapGroupsWithState operator grouped by dataSetKey.'
)

doc.add_heading('7.1 Operations', level=2)

add_table(
    ['Operation', 'Request ID', 'Action', 'Output'],
    [
        ['ADD', '1', 'Create synopsis via SynopsisFactory, store in state', 'None (or error)'],
        ['DATA', '(datapoint)', 'Iterate all active synopses, call add()', 'None'],
        ['ESTIMATE', '3', 'Call synopsis.estimate(), emit result', 'Estimation'],
        ['DELETE', '2', 'Remove synopsis from state', 'None'],
        ['TIMEOUT', '(automatic)', 'Evict all synopses for key, emit notices', 'Eviction notices'],
    ]
)

doc.add_heading('7.2 Event Ordering Within Micro-batches', level=2)
doc.add_paragraph(
    'Like the DataRouter, the SynopsisProcessor sorts events within each micro-batch '
    'to ensure correct ordering: ADD requests are processed first (so synopses exist '
    'before data arrives), then DATA events (so data is added before estimates), then '
    'ESTIMATE/DELETE requests. This is critical because Spark does not guarantee event '
    'ordering within a micro-batch.'
)

add_code("""// Process in order: ADDs -> DATA -> ESTIMATE/DELETE
for (InputEvent event : addRequests) handleRequest(...);
for (InputEvent event : dataEvents) handleData(...);
for (InputEvent event : otherRequests) handleRequest(...);""")

doc.add_heading('7.3 SynopsisProcessorState', level=2)
doc.add_paragraph(
    'The state holds a Map<Integer, Synopsis> mapping uid to synopsis instance. '
    'Before each state update, snapshotForSerialization() is called to ensure that '
    'synopsis objects with transient third-party fields are properly captured as byte arrays '
    'for Kryo serialization.'
)

add_code("""public class SynopsisProcessorState implements Serializable {
    private final Map<Integer, Synopsis> synopses = new HashMap<>();
    private Duration ttl;

    public void snapshotForSerialization() {
        for (Synopsis synopsis : synopses.values()) {
            synopsis.snapshotState();  // Capture transient fields
        }
    }
}""")

doc.add_heading('7.4 Safety Net TTL', level=2)
doc.add_paragraph(
    'Each key in the SynopsisProcessor has a configurable TTL (default: 7 days). '
    'If no events arrive for a key within this period, the state times out and all '
    'synopses for that key are evicted. Eviction notices are emitted to the output '
    'topic to inform downstream consumers. This prevents unbounded state growth '
    'from orphaned synopses that were never explicitly deleted.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 8. LAYER 4 — PATH SPLITTING
# ════════════════════════════════════════════════════════════════════

doc.add_heading('8. Layer 4 — Path Splitting', level=1)

doc.add_paragraph(
    'The PathSplitter routes estimations based on their parallelism level. '
    'This is a simple filter operation that separates the estimation stream into '
    'two paths:'
)

doc.add_paragraph('GREEN Path (noOfP == 1): Direct to output — no aggregation needed.', style='List Bullet')
doc.add_paragraph('PURPLE Path (noOfP > 1): Routed to the ReduceAggregator for merging.', style='List Bullet')

add_code("""public class PathSplitter {
    private final Dataset<Estimation> greenPath;
    private final Dataset<Estimation> purplePath;

    public PathSplitter(Dataset<Estimation> estimations) {
        this.greenPath  = estimations.filter(e -> e.getNoOfP() == 1);
        this.purplePath = estimations.filter(e -> e.getNoOfP() > 1);
    }
}""")

doc.add_paragraph(
    'This replaces Flink\'s SplitStream/OutputSelector pattern. In Spark, filtering '
    'creates two logical branches of the same streaming query — no data is duplicated.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 9. LAYER 5 — AGGREGATION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('9. Layer 5 — Aggregation (Purple Path)', level=1)

doc.add_paragraph(
    'The ReduceAggregator collects partial estimation results from multiple parallel '
    'synopsis partitions and merges them into a single final result. It is implemented '
    'as a flatMapGroupsWithState operator grouped by uid, so all partial results for '
    'the same synopsis instance land in the same operator instance.'
)

doc.add_heading('9.1 Aggregation Process', level=2)
doc.add_paragraph(
    'When the first partial arrives, a ReduceFunction is created based on the synopsisID. '
    'Subsequent partials are added to the reducer. When all noOfP partials have arrived, '
    'the reducer produces the final merged result and the state is cleared.'
)

doc.add_heading('9.2 Reduce Functions', level=2)
add_table(
    ['Function', 'Synopsis Types', 'Operation'],
    [
        ['SimpleSumFunction', 'CountMin (1), AMS (3), HLL (4)', 'Sums numerical partial estimations'],
        ['SimpleORFunction', 'BloomFilter (2)', 'Logical OR of boolean partial results'],
    ]
)

doc.add_heading('9.3 Timeout Handling', level=2)
doc.add_paragraph(
    'If not all partials arrive within the aggregation timeout (default: 5 minutes), '
    'the incomplete aggregation is discarded. This prevents state from growing indefinitely '
    'due to lost partials. The timeout is configurable via --aggregation-timeout.'
)

doc.add_heading('9.4 Improvement Over Flink', level=2)
doc.add_paragraph(
    'In the original Flink SDE, the final aggregation step (GReduceFlatMap) was forced to '
    'run with parallelism=1, creating a single-threaded bottleneck. In SDE_Spark, the '
    'ReduceAggregator runs with full parallelism because Spark\'s groupByKey(uid) naturally '
    'distributes different uids across partitions. Only partials for the same uid converge '
    'to the same operator instance.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 10. LAYER 6 — OUTPUT
# ════════════════════════════════════════════════════════════════════

doc.add_heading('10. Layer 6 — Output', level=1)

doc.add_paragraph(
    'KafkaOutputLayer serializes Estimation objects to JSON and writes them to the Kafka '
    'output topic. It replaces the FlinkKafkaProducer and addSink() from the Flink Run.java.'
)

add_code("""public StreamingQuery write(Dataset<Estimation> estimations, String queryName) {
    Dataset<String> jsonOutput = estimations.map(estimation -> {
        ObjectMapper mapper = new ObjectMapper();
        return mapper.writeValueAsString(estimation);
    }, Encoders.STRING());

    return jsonOutput.toDF("value")
        .selectExpr("value AS key", "value")
        .writeStream()
        .format("kafka")
        .option("kafka.bootstrap.servers", config.getKafkaBrokers())
        .option("topic", config.getOutputTopic())
        .option("kafka.acks", config.getKafkaProducerAcks())
        .option("kafka.compression.type", config.getKafkaProducerCompression())
        .option("kafka.enable.idempotence", ...)
        .option("checkpointLocation", ...)
        .trigger(Trigger.ProcessingTime(config.getTriggerInterval()))
        .start();
}""")

doc.add_heading('10.1 Producer Tuning', level=2)
add_table(
    ['Setting', 'Default', 'Purpose'],
    [
        ['kafka.acks', '"all"', 'Wait for all replicas to acknowledge (durability)'],
        ['kafka.compression.type', '"lz4"', 'LZ4 compression for reduced network usage'],
        ['kafka.enable.idempotence', 'true', 'Exactly-once delivery within Kafka'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 11. SYNOPSIS ALGORITHMS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('11. Synopsis Algorithms', level=1)

doc.add_paragraph(
    'SDE_Spark implements four probabilistic data structure algorithms, each extending '
    'the abstract Synopsis base class. The SynopsisFactory creates instances based on '
    'the synopsisID field in the Request message.'
)

add_table(
    ['ID', 'Algorithm', 'Use Case', 'Key Parameters'],
    [
        ['1', 'CountMin Sketch', 'Frequency estimation (e.g., total price per stock)', 'epsilon, delta, seed'],
        ['2', 'Bloom Filter', 'Membership testing (e.g., "has stock X been seen?")', 'expected items, FP rate'],
        ['3', 'AMS Sketch', 'Frequency moment estimation', 'buckets, depth'],
        ['4', 'HyperLogLog', 'Cardinality estimation (e.g., distinct stocks)', 'relative std dev'],
    ]
)

doc.add_heading('11.1 Synopsis Base Class', level=2)
add_code("""public abstract class Synopsis implements Serializable {
    protected int synopsisID;       // Unique instance ID
    protected String keyIndex;      // Field name for keys (e.g., "StockID")
    protected String valueIndex;    // Field name for values (e.g., "price")

    public abstract void add(JsonNode values);
    public abstract Object estimate(Object key);
    public abstract Estimation estimate(Request rq);
    public abstract Synopsis merge(Synopsis other);

    // Hook for Kryo serialization of transient fields
    public void snapshotState() { /* no-op by default */ }
}""")

doc.add_heading('11.2 CountMin Sketch (ID=1)', level=2)
doc.add_paragraph(
    'Wraps a custom CM implementation with configurable epsilon (error bound) and '
    'delta (confidence). Supports add(key, value) for accumulating values per key '
    'and estimateCount(key) for querying. The CM class uses a 2D array of counters '
    'with multiple hash functions (depth = ceil(ln(1/delta)), width = ceil(e/epsilon)).'
)
doc.add_paragraph(
    'Parameters: [keyField, valueField, operationMode, epsilon, delta, seed]'
)

doc.add_heading('11.3 Bloom Filter (ID=2)', level=2)
doc.add_paragraph(
    'Provides probabilistic membership testing. Uses the Guava-style BloomFilter from '
    'the stream-lib library. The snapshotState()/ensureBloomFilter() pattern handles '
    'Kryo serialization by converting the transient BloomFilter to/from a byte array.'
)
doc.add_paragraph(
    'Parameters: [keyField, valueField, operationMode, expectedInsertions, numHashFunctions]'
)

doc.add_heading('11.4 AMS Sketch (ID=3)', level=2)
doc.add_paragraph(
    'Estimates frequency moments of the data stream using the Alon-Matias-Szegedy algorithm. '
    'Implemented via the streaminer library. Uses reflection-based snapshot/restore for '
    'serialization of transient internal state.'
)
doc.add_paragraph(
    'Parameters: [keyField, valueField, operationMode, buckets, depth]'
)

doc.add_heading('11.5 HyperLogLog (ID=4)', level=2)
doc.add_paragraph(
    'Estimates the number of distinct elements in the stream. Uses the HyperLogLog '
    'implementation from stream-lib (Clearspring). Serializes internal registers as a '
    'byte array for Kryo compatibility.'
)
doc.add_paragraph(
    'Parameters: [keyField, valueField, operationMode, relativeStdDev]'
)

doc.add_heading('11.6 Kryo Serialization Pattern', level=2)
doc.add_paragraph(
    'Several synopsis implementations wrap third-party objects that are not natively '
    'serializable by Kryo. The solution is the snapshot/restore pattern: before '
    'state serialization, snapshotState() converts transient objects to byte arrays. '
    'On first use after deserialization, the byte array is used to reconstruct the '
    'transient object. This is critical for Spark\'s state checkpointing.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 12. CONFIGURATION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('12. Configuration & Deployment', level=1)

doc.add_paragraph(
    'SDEConfig is the centralized configuration class. It supports two-pass loading: '
    'first from an optional properties file (--config flag), then from CLI arguments '
    'that override file values.'
)

doc.add_heading('12.1 Configuration Parameters', level=2)

add_table(
    ['Parameter', 'CLI Flag', 'Default', 'Description'],
    [
        ['Data topic', '--data-topic', 'data_topic', 'Kafka topic for data events'],
        ['Request topic', '--request-topic', 'request_topic', 'Kafka topic for requests'],
        ['Output topic', '--output-topic', 'estimation_topic', 'Kafka topic for estimations'],
        ['Kafka brokers', '--kafka-brokers', 'localhost:9092', 'Bootstrap servers'],
        ['Trigger interval', '--trigger-interval', '1 second', 'Micro-batch trigger'],
        ['Checkpoint', '--checkpoint-location', 'hdfs:///sde/checkpoints', 'State checkpoint path'],
        ['Safety TTL', '--safety-net-ttl', '7 days', 'Synopsis state expiration'],
        ['Group ID', '--kafka-group-id', 'sde-spark', 'Kafka consumer group'],
        ['Offsets', '--starting-offsets', 'earliest', 'Initial offset strategy'],
        ['Routing TTL', '--routing-state-ttl', '1 day', 'Routing state expiration'],
        ['Aggregation timeout', '--aggregation-timeout', '5 minutes', 'Purple path timeout'],
        ['Fail on data loss', '--fail-on-data-loss', 'false', 'Strict offset checking'],
    ]
)

doc.add_heading('12.2 Kafka Security', level=2)
doc.add_paragraph(
    'SDE_Spark supports SSL/TLS encryption and SASL authentication for Kafka connections. '
    'Security options are applied to both consumers (Layer 1) and producers (Layer 6) '
    'via the getKafkaSecurityOptions() helper.'
)
add_table(
    ['Parameter', 'CLI Flag', 'Description'],
    [
        ['Security protocol', '--kafka-security-protocol', 'SASL_SSL, SSL, SASL_PLAINTEXT'],
        ['SASL mechanism', '--kafka-sasl-mechanism', 'PLAIN, SCRAM-SHA-256/512, GSSAPI'],
        ['JAAS config', '--kafka-sasl-jaas-config', 'Full JAAS configuration string'],
        ['Truststore location', '--kafka-ssl-truststore-location', 'Path to SSL truststore'],
        ['Truststore password', '--kafka-ssl-truststore-password', 'Truststore password'],
    ]
)

doc.add_heading('12.3 Deployment', level=2)
doc.add_paragraph(
    'SDE_Spark is packaged as a fat JAR using the Maven Shade plugin. It can be '
    'submitted to any Spark cluster:'
)
add_code("""# Local development
spark-submit --master local[4] \\
  --class infore.sde.spark.SDESparkApp \\
  target/sde-spark-1.0.0-SNAPSHOT.jar \\
  --kafka-brokers localhost:9092

# Cluster deployment
spark-submit --master yarn --deploy-mode cluster \\
  --num-executors 4 --executor-memory 4g \\
  --class infore.sde.spark.SDESparkApp \\
  sde-spark-1.0.0-SNAPSHOT.jar \\
  --config /etc/sde/production.properties""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 13. STATE MANAGEMENT
# ════════════════════════════════════════════════════════════════════

doc.add_heading('13. State Management & Serialization', level=1)

doc.add_paragraph(
    'SDE_Spark uses Spark\'s built-in HDFSBackedStateStore for state persistence. '
    'This was chosen over RocksDB after careful evaluation of trade-offs.'
)

doc.add_heading('13.1 State Store Choice', level=2)
add_table(
    ['Aspect', 'HDFSBackedStateStore (chosen)', 'RocksDB'],
    [
        ['Storage', 'In-memory with HDFS checkpoints', 'On-disk with in-memory cache'],
        ['Performance', 'Fast for small-medium state', 'Better for very large state'],
        ['Complexity', 'Zero configuration', 'Requires tuning'],
        ['Recovery', 'Full state reload from checkpoint', 'Incremental snapshots'],
        ['SDE fit', 'Good (synopses are compact)', 'Overkill for typical use'],
    ]
)

doc.add_heading('13.2 Stateful Operators', level=2)
doc.add_paragraph(
    'Three operators maintain state via flatMapGroupsWithState:'
)
add_table(
    ['Operator', 'State Class', 'Grouped By', 'Timeout'],
    [
        ['DataRouter', 'RoutingState', 'dataSetKey', '1 day (configurable)'],
        ['SynopsisProcessor', 'SynopsisProcessorState', 'dataSetKey', '7 days (safety net)'],
        ['ReduceAggregator', 'AggregationState', 'uid', '5 minutes (configurable)'],
    ]
)

doc.add_heading('13.3 Kryo Serialization', level=2)
doc.add_paragraph(
    'All state objects and message POJOs are serialized using Kryo (configured globally '
    'via spark.serializer). The key challenge was that several synopsis implementations '
    'wrap third-party objects with transient fields that Kryo cannot serialize directly. '
    'The solution is the snapshot/restore byte array pattern described in Section 11.6.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 14. MONITORING & METRICS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('14. Monitoring & Metrics', level=1)

doc.add_heading('14.1 Pipeline Metrics (Spark Accumulators)', level=2)
doc.add_paragraph(
    'PipelineMetrics provides 9 Spark LongAccumulators for real-time monitoring:'
)
add_table(
    ['Accumulator', 'Incremented By', 'Purpose'],
    [
        ['datapointsProcessed', 'SynopsisProcessor', 'Total data events processed'],
        ['requestsProcessed', 'SynopsisProcessor', 'Total request events processed'],
        ['synopsesCreated', 'SynopsisProcessor (ADD)', 'Synopsis instances created'],
        ['synopsesDeleted', 'SynopsisProcessor (DELETE)', 'Synopsis instances deleted'],
        ['estimationsEmitted', 'SynopsisProcessor (ESTIMATE)', 'Estimation results produced'],
        ['aggregationsCompleted', 'ReduceAggregator', 'Purple path merges completed'],
        ['parseErrors', 'KafkaIngestionLayer', 'JSON deserialization failures'],
        ['routingFanOuts', 'DataRouter', 'Data fan-out operations'],
        ['stateTimeouts', 'SynopsisProcessor', 'State timeout evictions'],
    ]
)

doc.add_heading('14.2 ThroughputListener', level=2)
doc.add_paragraph(
    'A StreamingQueryListener that captures per-batch metrics and writes them to CSV '
    'for post-experiment analysis. Metrics per batch:'
)
doc.add_paragraph('timestamp — wall-clock time of batch completion', style='List Bullet')
doc.add_paragraph('batch_id — sequential batch number', style='List Bullet')
doc.add_paragraph('num_input_rows — rows read from Kafka in this batch', style='List Bullet')
doc.add_paragraph('input_rows_per_sec — Spark-reported input rate', style='List Bullet')
doc.add_paragraph('processed_rows_per_sec — Spark-reported processing rate', style='List Bullet')
doc.add_paragraph('batch_duration_ms — total time to process this batch', style='List Bullet')

doc.add_paragraph(
    'On query termination, the listener writes a summary file with aggregate statistics '
    '(total rows, average throughput, peak throughput, average batch duration).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 15. E2E VALIDATED FLOWS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('15. End-to-End Validated Flows', level=1)

doc.add_paragraph(
    'Both the GREEN and PURPLE paths have been validated end-to-end using the '
    'EndToEndTestApp and EndToEndTestProducer. The test sends known data through '
    'the full pipeline and verifies the estimation output on the Kafka output topic.'
)

doc.add_heading('15.1 GREEN Path Validation (noOfP=1)', level=2)

doc.add_paragraph('Test scenario: CountMin Sketch with uid=10, querying stock "AAPL"')
doc.add_paragraph('Input data:')
add_table(
    ['Stream', 'StockID', 'Price'],
    [
        ['EURUSD', 'AAPL', '182.50'],
        ['EURUSD', 'AAPL', '183.00'],
        ['EURUSD', 'GOOG', '141.20'],
        ['EURUSD', 'AAPL', '184.00'],
        ['EURUSD', 'MSFT', '378.90'],
    ]
)

doc.add_paragraph(
    'ESTIMATE("AAPL") result: 549.0 (= 182.50 + 183.00 + 184.00)'
)
doc.add_paragraph(
    'Flow: Kafka -> Layer 1 (parse) -> Layer 2 (pass-through, noOfP=1) -> '
    'Layer 3 (CountMin add + estimate) -> Layer 4 (GREEN filter) -> '
    'Layer 6 (output to estimation_topic)'
)

doc.add_heading('15.2 PURPLE Path Validation (noOfP=2)', level=2)

doc.add_paragraph('Test scenario: CountMin Sketch with uid=50, noOfP=2, querying stock "AAPL"')
doc.add_paragraph('Data distribution after hash-based routing:')
add_table(
    ['Partition', 'Data Received', 'AAPL Partial'],
    [
        ['Forex_2_KEYED_0', 'AAPL 185.00, TSLA 248.50', '185.0'],
        ['Forex_2_KEYED_1', 'AAPL 186.00, AAPL 187.00, NVDA 920.00, GOOG 142.00', '373.0'],
    ]
)

doc.add_paragraph(
    'After aggregation: 185.0 + 373.0 = 558.0'
)
doc.add_paragraph(
    'Flow: Layer 2 fans out ADD to both partitions -> Layer 3 creates '
    'independent CountMin in each partition -> Layer 3 adds data per partition -> '
    'Layer 3 produces partial estimates -> Layer 4 PURPLE filter -> '
    'Layer 5 ReduceAggregator merges via SimpleSumFunction -> Layer 6 output'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 16. UNIT TESTING
# ════════════════════════════════════════════════════════════════════

doc.add_heading('16. Unit Testing', level=1)

doc.add_paragraph(
    'SDE_Spark includes 12 test classes with 51 unit tests covering the core components. '
    'Tests use JUnit 5 and AssertJ for assertions. All tests pass successfully.'
)

add_table(
    ['Test Class', 'Tests', 'Coverage Area'],
    [
        ['SynopsisFactoryTest', '7', 'Factory creation for all 4 types, error handling for invalid/null params'],
        ['CountMinTest', '3', 'add() + estimate(), estimate via Request, count tracking'],
        ['SynopsisSerializationTest', '4', 'Serialization round-trip for all 4 synopsis types'],
        ['RoutingStateTest', '3', 'Registration, parallelism tracking, unregistration'],
        ['RequestRouterTest', '6', 'Fan-out for noOfP=1 and noOfP>1, operation filtering'],
        ['InputEventTest (routing)', '2', 'Data and request wrapping in InputEvent'],
        ['InputEventTest (processing)', '2', 'InputEvent discriminator type checking'],
        ['SDEConfigTest', '7', 'Default values, CLI overrides, file loading, security options'],
        ['SimpleSumFunctionTest', '3', 'Sum aggregation with 2 and 3 partials, single partial'],
        ['ReduceFunctionTest', '4', 'SimpleSumFunction + SimpleORFunction integration'],
        ['MessageSerializationTest', '5', 'JSON round-trip for Datapoint, Request, Estimation (Kafka compat)'],
        ['SynopsisProcessorTest', '5', 'ADD + ESTIMATE for all 4 algorithms'],
    ]
)

doc.add_heading('16.1 Testing Strategy', level=2)
doc.add_paragraph(
    'The testing strategy focuses on three areas:'
)
doc.add_paragraph(
    'Algorithm correctness: Each synopsis type is tested for correct add/estimate behavior '
    'with known inputs and expected outputs.',
    style='List Bullet'
)
doc.add_paragraph(
    'Serialization integrity: All synopses and messages are tested for round-trip '
    'serialization (Java Serialization and Kryo) to ensure state persistence works correctly.',
    style='List Bullet'
)
doc.add_paragraph(
    'Component behavior: Routing state, configuration loading, aggregation functions, and '
    'input event handling are tested in isolation.',
    style='List Bullet'
)

doc.add_heading('16.2 Running Tests', level=2)
add_code("""# Run all tests
mvn test

# Run a specific test class
mvn test -Dtest=CountMinTest

# Run with verbose output
mvn test -Dsurefire.useFile=false""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 17. EXPERIMENTAL EVALUATION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('17. Experimental Evaluation', level=1)

doc.add_paragraph(
    'Following the methodology of Kontaxakis (2020, Chapter 5), we evaluated '
    'the SDE_Spark pipeline\'s throughput scaling behavior. This section presents '
    'preliminary results from Experiment B: Throughput vs Ingestion Rate.'
)

doc.add_heading('17.1 Experiment Infrastructure', level=2)

doc.add_paragraph(
    'Two components were built specifically for experimental evaluation:'
)

p = doc.add_paragraph()
run = p.add_run('LoadTestProducer: ')
run.bold = True
p.add_run(
    'A configurable Kafka producer that generates synthetic stock market data at '
    'a controlled, steady rate. Supports configurable parameters: target rate, duration, '
    'number of streams, number of synopses, parallelism level, warmup period, and '
    'estimate interval. Operates in 4 phases: register synopses, warmup, measured '
    'sending with periodic estimates, and cleanup.'
)

p = doc.add_paragraph()
run = p.add_run('ThroughputListener: ')
run.bold = True
p.add_run(
    'A StreamingQueryListener that captures per-batch metrics (input rows, input rate, '
    'processed rate, batch duration) and writes them to CSV for analysis. Also produces '
    'aggregate statistics on query termination.'
)

doc.add_heading('17.2 Test Setup', level=2)

add_table(
    ['Component', 'Specification'],
    [
        ['CPU', 'AMD Ryzen 7 5800H (8 cores / 16 threads)'],
        ['RAM', '16 GB'],
        ['Storage', '512 GB NVMe SSD'],
        ['OS', 'Windows 11'],
        ['Spark', '3.5.3 (Structured Streaming, local[4] mode)'],
        ['Kafka', '7.5.3 (single broker, Docker)'],
        ['Java', '17, Kryo serialization'],
        ['Shuffle partitions', '4'],
        ['Trigger interval', '2 seconds'],
        ['Driver memory', '2 GB'],
    ]
)

doc.add_heading('17.3 Experiment B: Throughput vs Ingestion Rate', level=2)

doc.add_paragraph(
    'Three runs were executed at increasing target ingestion rates: 100, 500, and 1000 '
    'messages per second. Each run lasted 60 seconds (plus 10 seconds warmup). '
    'A single CountMin Sketch synopsis (epsilon=0.002, delta=0.01) was registered on '
    '10 streams (GREEN path, noOfP=1). 50 stock symbols with random prices [50-1000] '
    'were used as input data.'
)

doc.add_heading('17.4 Results', level=2)

add_table(
    ['Target Rate', 'Avg Input (rows/sec)', 'Avg Processed (rows/sec)', 'Peak (rows/sec)', 'Avg Batch (ms)', 'Total Rows'],
    [
        ['100', '87.3', '86.4', '107.7', '6,805', '6,903'],
        ['500', '454.2', '447.7', '524.3', '6,786', '36,302'],
        ['1,000', '956.8', '959.6', '1,035.5', '6,356', '67,101'],
    ]
)

doc.add_heading('17.5 Key Findings', level=2)

p = doc.add_paragraph()
run = p.add_run('a) Linear Throughput Scaling: ')
run.bold = True
p.add_run(
    'The pipeline sustains all tested ingestion rates. When the input rate increases '
    'from 100 to 1,000 msg/sec (10x), the processed rate increases proportionally '
    '(86 to 960 rows/sec, ~11x). At no point did the pipeline fall behind.'
)

p = doc.add_paragraph()
run = p.add_run('b) Constant Batch Duration: ')
run.bold = True
p.add_run(
    'Average batch duration remains ~6.5 seconds regardless of the ingestion rate. '
    'This indicates the bottleneck is per-batch overhead (state serialization, Kafka '
    'polling, task scheduling), not synopsis computation.'
)

p = doc.add_paragraph()
run = p.add_run('c) Synopsis Correctness: ')
run.bold = True
p.add_run(
    'The CountMin sketch correctly accumulated price values and returned a valid '
    'estimation (60,302.0 for stock "ABBV"). No data corruption or parsing errors '
    'were observed during the measurement phase.'
)

p = doc.add_paragraph()
run = p.add_run('d) Comparison with Kontaxakis (2020): ')
run.bold = True
p.add_run(
    'These results mirror the pattern shown in Figure 24 of the original thesis: '
    'near-linear throughput scaling with increasing ingestion rate. The SDE_Spark '
    'implementation preserves the scalability characteristics of the Flink-based SDE.'
)

doc.add_heading('17.6 Limitations', level=2)
limitations = [
    'Single-node execution: all Spark executors share the same CPU and memory. '
    'Multi-node cluster experiments are needed to evaluate horizontal scalability.',
    'Limited stream count (10 vs 50/500/5000 in Kontaxakis) due to 16 GB RAM constraint.',
    'Only GREEN path tested (noOfP=1). PURPLE path requires separate evaluation.',
    'No direct comparison against SDE_Flink on identical hardware.',
]
for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

doc.add_heading('17.7 Planned Full-Scale Experiments', level=2)
add_table(
    ['Experiment', 'Variable', 'Range', 'Kontaxakis Equivalent'],
    [
        ['A: Workers', 'Spark executors', '2, 4, 6, 8, 10', 'Figure 23'],
        ['B: Ingestion Rate', 'Messages/sec', '1K, 2K, 5K, 10K', 'Figure 24'],
        ['C: Streams', 'Number of streams', '50, 500, 5000', 'Figure 25'],
        ['D: Communication', 'Shuffle bytes', '50, 500, 5000 streams', 'Figure 26'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 18. COMPARISON WITH SDE_FLINK
# ════════════════════════════════════════════════════════════════════

doc.add_heading('18. Comparison with SDE_Flink', level=1)

doc.add_paragraph(
    'The following table summarizes the key architectural differences between '
    'the original Flink implementation and the Spark re-implementation.'
)

add_table(
    ['Aspect', 'SDE_Flink', 'SDE_Spark'],
    [
        ['Framework', 'Apache Flink 1.9.3', 'Apache Spark 3.5.3'],
        ['Processing model', 'True streaming (event-at-a-time)', 'Micro-batch (Structured Streaming)'],
        ['Stateful API', 'CoProcessFunction', 'flatMapGroupsWithState'],
        ['Two-input streams', 'ConnectedStreams (native)', 'Union + InputEvent discriminator'],
        ['State backend', 'RocksDB (Flink-managed)', 'HDFSBackedStateStore (Spark-managed)'],
        ['Serialization', 'Java Serialization', 'Kryo (with snapshot/restore pattern)'],
        ['Routing', 'Separate RqRouter + DataRouter', 'Unified DataRouter (single operator)'],
        ['Aggregation bottleneck', 'GReduce parallelism=1', 'Full parallelism via groupByKey'],
        ['Synopsis IDs', '1,2,3,7 (from 29-ID registry)', '1,2,3,4 (clean sequential)'],
        ['Event ordering', 'Guaranteed per-key', 'Explicit sorting within micro-batch'],
        ['Checkpointing', 'Flink checkpoints', 'Spark checkpoints to HDFS/S3'],
        ['Kafka contract', 'Same JSON schema', 'Same JSON schema (compatible)'],
    ]
)

doc.add_heading('18.1 Key Improvements in SDE_Spark', level=2)

p = doc.add_paragraph()
run = p.add_run('Unified DataRouter: ')
run.bold = True
p.add_run(
    'Combining request and data routing into a single stateful operator eliminates '
    'the potential for state inconsistency between separate operators.'
)

p = doc.add_paragraph()
run = p.add_run('No aggregation bottleneck: ')
run.bold = True
p.add_run(
    'Spark\'s groupByKey naturally distributes aggregation across partitions, '
    'eliminating the GReduceFlatMap.setParallelism(1) bottleneck in Flink.'
)

p = doc.add_paragraph()
run = p.add_run('Explicit event ordering: ')
run.bold = True
p.add_run(
    'Sorting events within each micro-batch makes ordering explicit and deterministic, '
    'rather than relying on implicit stream ordering guarantees.'
)

p = doc.add_paragraph()
run = p.add_run('Modern technology stack: ')
run.bold = True
p.add_run(
    'Spark 3.5.3 with Java 17 benefits from years of Spark improvements in structured '
    'streaming, state management, and Kafka integration.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 19. CONCLUSIONS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('19. Conclusions & Future Work', level=1)

doc.add_heading('19.1 Conclusions', level=2)
doc.add_paragraph(
    'SDE_Spark successfully re-implements the Synopsis Data Engine on Apache Spark '
    'Structured Streaming while preserving full compatibility with the original Flink '
    'implementation\'s Kafka API contract. The key contributions are:'
)

contributions = [
    'A complete 6-layer pipeline implementation that handles both GREEN (single-partition) '
    'and PURPLE (multi-partition with aggregation) processing paths.',
    'A novel unified DataRouter that combines request and data routing in a single '
    'stateful operator, addressing the two-input stream challenge in Spark.',
    'Implementation of all 4 synopsis algorithms (CountMin, Bloom Filter, AMS, HyperLogLog) '
    'with a Kryo-compatible serialization pattern.',
    'Production-ready configuration supporting properties files, CLI overrides, Kafka security '
    '(SASL/SSL), and tunable timeouts.',
    'Comprehensive testing: 51 unit tests and validated end-to-end flows for both paths.',
    'Preliminary experimental evaluation showing linear throughput scaling, consistent with '
    'Kontaxakis (2020) results.',
]
for c in contributions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('19.2 Future Work', level=2)
future = [
    'Full-scale cluster experiments (Experiments A-D) on a multi-node Spark cluster.',
    'Direct SDE_Spark vs SDE_Flink performance comparison on identical hardware.',
    'Additional synopsis algorithms (e.g., Coresets, Wavelets) to expand the SDE repertoire.',
    'Dynamic scaling: automatically adjusting parallelism based on data volume.',
    'Real-time monitoring dashboard integrating PipelineMetrics accumulators.',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'SDE-Spark-Implementation-Report.docx'
)
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Total sections: 19")
